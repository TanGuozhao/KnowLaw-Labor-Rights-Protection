"""民事起诉状 Word 模板：段落 + {{占位符}}，由 routes 填充。"""
from __future__ import annotations

from pathlib import Path

from docx import Document

CIVIL_COMPLAINT_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent / "docx" / "civilComplaint.docx"
)


def ensure_civil_complaint_template_exists() -> Path:
    """若模板不存在则生成一份默认段落占位符文档。"""
    path = CIVIL_COMPLAINT_TEMPLATE_PATH
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        return path
    doc = Document()
    doc.add_paragraph("民事起诉状")
    doc.add_paragraph("")
    doc.add_paragraph("原告：{{plaintiff_block}}")
    doc.add_paragraph("")
    doc.add_paragraph("被告：{{defendant_block}}")
    doc.add_paragraph("")
    doc.add_paragraph("案由：{{case_cause}}")
    doc.add_paragraph("")
    doc.add_paragraph("诉讼请求")
    doc.add_paragraph("{{claims}}")
    doc.add_paragraph("")
    doc.add_paragraph("事实与理由")
    doc.add_paragraph("{{facts}}")
    doc.add_paragraph("")
    doc.add_paragraph("证据和证据来源（如有）")
    doc.add_paragraph("{{evidence_list}}")
    doc.add_paragraph("")
    doc.add_paragraph("此致")
    doc.add_paragraph("{{court_name}}")
    doc.add_paragraph("")
    doc.add_paragraph("具状人：{{plaintiff_name}}（签字）")
    doc.add_paragraph("{{date_text}}")
    doc.save(str(path))
    return path


def replace_placeholders_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    text = paragraph.text
    if "{{" not in text:
        return
    for key, value in mapping.items():
        text = text.replace("{{" + key + "}}", value)
    paragraph.text = text


def replace_placeholders_in_document(doc: Document, mapping: dict[str, str]) -> None:
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, mapping)


def format_plaintiff_block(payload: dict) -> str:
    parts: list[str] = []
    name = str(payload.get("plaintiff_name") or "").strip()
    if name:
        parts.append(f"姓名：{name}")
    g = str(payload.get("plaintiff_gender") or "").strip()
    if g:
        parts.append(f"性别：{g}")
    eth = str(payload.get("plaintiff_ethnicity") or "").strip()
    if eth:
        parts.append(f"民族：{eth}")
    birth = str(payload.get("plaintiff_birth") or "").strip()
    if birth:
        parts.append(f"出生日期：{birth}")
    addr = str(payload.get("plaintiff_address") or "").strip()
    if addr:
        parts.append(f"住址：{addr}")
    idn = str(payload.get("plaintiff_id_number") or "").strip()
    if idn:
        parts.append(f"公民身份号码：{idn}")
    phone = str(payload.get("plaintiff_phone") or "").strip()
    if phone:
        parts.append(f"联系电话：{phone}")
    if not parts:
        return "（请填写原告信息）"
    return "，".join(parts)


def format_defendant_block(payload: dict) -> str:
    org_rep = str(payload.get("defendant_legal_representative") or "").strip()
    name = str(payload.get("defendant_name") or "").strip()
    addr = str(payload.get("defendant_address") or "").strip()
    phone = str(payload.get("defendant_phone") or "").strip()
    if org_rep:
        parts: list[str] = []
        if name:
            parts.append(f"单位名称：{name}")
        if addr:
            parts.append(f"住所地：{addr}")
        parts.append(f"法定代表人/主要负责人：{org_rep}")
        if phone:
            parts.append(f"联系电话：{phone}")
        return "，".join(parts) if parts else "（请填写被告单位信息）"
    parts = []
    if name:
        parts.append(f"姓名：{name}")
    if addr:
        parts.append(f"住址：{addr}")
    if phone:
        parts.append(f"联系电话：{phone}")
    if not parts:
        return "（请填写被告信息）"
    return "，".join(parts)


def build_civil_complaint_mapping(payload: dict) -> dict[str, str]:
    """前端 JSON -> 模板占位符（值均为字符串）。"""
    return {
        "plaintiff_block": format_plaintiff_block(payload),
        "defendant_block": format_defendant_block(payload),
        "case_cause": str(payload.get("case_cause") or "").strip() or "（案由）",
        "claims": str(payload.get("claims") or "").strip() or "（诉讼请求）",
        "facts": str(payload.get("facts") or "").strip() or "（事实与理由）",
        "evidence_list": str(payload.get("evidence_list") or "").strip() or "（无）",
        "court_name": str(payload.get("court_name") or "").strip() or "（受诉人民法院全称）",
        "plaintiff_name": str(payload.get("plaintiff_name") or "").strip() or "（原告姓名）",
        "date_text": str(payload.get("date_text") or "").strip() or "（日期）",
    }
