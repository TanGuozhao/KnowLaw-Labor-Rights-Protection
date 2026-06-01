# -*- coding: utf-8 -*-
"""
将用户提供的范文 Word 转为可程序填充的 {{mustache}} 模板：
  - laborArbitrationApplicationForm.docx
  - laborDisputeMediationApplicationForm.docx
  - applicationForEnforcementDocument.docx → 生成 applicationForEnforcementDocumentForApp.docx

运行：python backend/document_writing/patch_labor_word_templates.py （仓库根目录）
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "docx"


def patch_arbitration(path: Path) -> None:
    d = Document(str(path))
    ps = d.paragraphs  # 与范文原始段落顺序一致（劳动人事争议仲裁申请书）
    ps[1].text = "劳动人事争议仲裁申请书"
    ps[3].text = "申请人：{{applicant_block}}"
    ps[4].text = "委托代理人：{{agent_block}}"
    ps[6].text = "被申请人：{{respondent_block}}"
    ps[7].text = ""
    ps[9].text = "请求事项：（请求要明确，涉及金额要有具体的计算标准和过程，如计算过程复杂，可作为附件提交）"
    ps[10].text = "{{claims}}"
    ps[11].text = ""
    ps[13].text = "事实与理由：（简明扼要写清楚入职时间、争议时间和内容、离职时间等）"
    ps[14].text = "{{facts}}"
    ps[15].text = "证据和证据来源（如有）\n{{evidence_list}}"
    ps[16].text = "此致"
    ps[17].text = "{{arbitration_commission}}"
    ps[18].text = ""
    ps[19].text = "申请人：{{applicant_name}}（必须本人签字）"
    ps[20].text = "{{date_text}}"
    ps[21].text = "（应为提交当天的日期）"
    d.save(str(path))


def patch_mediation(path: Path) -> None:
    d = Document(str(path))
    ps = d.paragraphs
    # 原0–13 条替换为占位符版式（保留调解申请书结构）
    body = [
        "申请人：{{applicant_block}}",
        "被申请人：{{respondent_block}}",
        "事由：{{mediation_preamble}}",
        "调解请求：",
        "{{claims}}",
        "事实与理由：",
        "{{facts}}",
        "证据和证据来源（如有）\n{{evidence_list}}",
        "为此，向{{mediation_org}}申请调解，请依法调解。",
        "申请人：{{applicant_name}}（签名或盖章）",
        "{{date_text}}",
    ]
    for i, text in enumerate(body):
        if i < len(ps):
            ps[i].text = text
        else:
            d.add_paragraph(text)
    # 清空多余旧段落（若原文件更长）
    for j in range(len(body), len(ps)):
        ps[j].text = ""
    d.save(str(path))


def patch_enforcement(source: Path, dest: Path) -> None:
    """法院样式「申请执行书」范文嵌入占位符；写入 dest（避免源文件被 Word 占用无法保存）。"""
    d = Document(str(source))
    ps = d.paragraphs
    ps[2].text = "申请执行人：{{applicant_block}}"
    ps[3].text = "法定代理人/指定代理人：{{legal_representative_line}}"
    ps[4].text = "委托诉讼代理人：{{entrusted_agent_line}}"
    ps[5].text = "被执行人：{{respondent_block}}"
    ps[6].text = ""
    # ps[7] 保持说明原文
    ps[8].text = "{{enforcement_opening_paragraph}}"
    ps[9].text = "请求事项"
    ps[10].text = "{{requests}}"
    ps[11].text = "此致"
    ps[12].text = "{{court_name}}"
    ps[14].text = "{{attachment_line}}"
    ps[15].text = "申请执行人(签名或盖章)"
    ps[16].text = "{{date_text}}"
    d.save(str(dest))


def main() -> None:
    arb = DOCX / "laborArbitrationApplicationForm.docx"
    med = DOCX / "laborDisputeMediationApplicationForm.docx"
    enf_src = DOCX / "applicationForEnforcementDocument.docx"
    enf_app = DOCX / "applicationForEnforcementDocumentForApp.docx"
    if not arb.exists():
        print("Missing:", arb, file=sys.stderr)
        sys.exit(1)
    if not med.exists():
        print("Missing:", med, file=sys.stderr)
        sys.exit(1)
    patch_arbitration(arb)
    patch_mediation(med)
    if enf_src.exists():
        patch_enforcement(enf_src, enf_app)
        print("Patched:", arb.name, med.name, "->", enf_app.name)
    else:
        print("Patched:", arb.name, med.name, "(skip enforcement:", enf_src.name, "missing)")


if __name__ == "__main__":
    main()
