import json
import re
import sys
from pathlib import Path

from docx import Document

BACKEND_ROOT = Path(__file__).resolve().parent.parent
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from chat_service import chat_completion


def _parse_json(text: str) -> dict:
    raw = str(text or "").strip()
    if not raw:
        raise RuntimeError("LLM returned empty response")
    try:
        data = json.loads(raw)
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass
    m = re.search(r"\{[\s\S]*\}", raw)
    if not m:
        raise RuntimeError("LLM did not return parseable JSON")
    data = json.loads(m.group(0))
    if not isinstance(data, dict):
        raise RuntimeError("LLM JSON root is not an object")
    return data


def _fill_cell(table, row: int, col: int, value: str) -> None:
    cell = table.cell(row, col)
    if not cell.paragraphs:
        cell.text = value
        return
    p = cell.paragraphs[0]
    p.clear()
    p.add_run(value)


def main() -> None:
    schema = [
        "complainant_name",
        "complainant_gender",
        "complainant_mobile_phone",
        "complainant_id_number",
        "complainant_mailing_address",
        "complainant_landline_phone",
        "complainant_postal_code",
        "respondent_name",
        "respondent_legal_representative",
        "respondent_contact_name",
        "respondent_contact_job_title",
        "respondent_registered_address",
        "respondent_business_address",
        "respondent_contact_phone",
        "respondent_postal_code",
        "claim_requests",
        "facts_and_reasons",
    ]

    system = (
        "Generate realistic fictional data for a labor security inspection complaint form in China. "
        "Return JSON only. Include all keys and ensure values are strings."
    )
    user = "Keys: " + ",".join(schema)
    text = chat_completion(
        [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ]
    )
    data = _parse_json(text)
    fields = {k: str(data.get(k, "") or "").strip() for k in schema}

    template = Path(r"f:\LabelHelp\backend\docx\laborSecurityInspectionComplaintForm.docx")
    out_dir = Path(r"f:\LabelHelp\backend\docx\output")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "laborSecurityInspectionComplaintForm.sample.llm.docx"

    doc = Document(str(template))
    table = doc.tables[0]

    _fill_cell(table, 0, 2, fields["complainant_name"])
    _fill_cell(table, 0, 4, fields["complainant_gender"])
    _fill_cell(table, 0, 6, fields["complainant_mobile_phone"])
    _fill_cell(table, 1, 2, fields["complainant_id_number"])
    _fill_cell(table, 1, 6, fields["complainant_landline_phone"])
    _fill_cell(table, 2, 2, fields["complainant_mailing_address"])
    _fill_cell(table, 2, 6, fields["complainant_postal_code"])
    _fill_cell(table, 3, 2, fields["respondent_name"])
    _fill_cell(table, 3, 6, fields["respondent_registered_address"])
    _fill_cell(table, 4, 6, fields["respondent_business_address"])
    _fill_cell(table, 5, 2, fields["respondent_legal_representative"])
    _fill_cell(table, 6, 2, fields["respondent_contact_name"])
    _fill_cell(table, 5, 4, fields["respondent_contact_job_title"])
    _fill_cell(table, 6, 4, fields["respondent_contact_job_title"])
    _fill_cell(table, 5, 6, fields["respondent_contact_phone"])
    _fill_cell(table, 6, 6, fields["respondent_postal_code"])
    _fill_cell(table, 7, 0, fields["claim_requests"])
    _fill_cell(table, 8, 0, fields["facts_and_reasons"])

    doc.save(str(out_path))
    print(json.dumps({"output": str(out_path), "fields": fields}, ensure_ascii=False))


if __name__ == "__main__":
    main()
