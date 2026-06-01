"""证据材料清单：程序生成表格版 Word（与常见法院样式一致）。"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

# 版式：标题仿宋_GB2312 一号加粗；其余仿宋_GB2312 小三；表头行加粗。
FONT_FANGSONG_GB2312 = "仿宋_GB2312"
PT_YI_HAO = 26  # 一号
PT_XIAO_SAN = 15  # 小三

EVIDENCE_LIST_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent / "docx" / "evidenceListMaterials.docx"
)


def _fmt_run(run, size_pt: float, *, bold: bool = False) -> None:
    """统一设置中英文字体为仿宋_GB2312（Word 中文字体依赖 w:eastAsia）。"""
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = FONT_FANGSONG_GB2312
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_FANGSONG_GB2312)
    r_fonts.set(qn("w:hAnsi"), FONT_FANGSONG_GB2312)
    r_fonts.set(qn("w:eastAsia"), FONT_FANGSONG_GB2312)
    r_fonts.set(qn("w:cs"), FONT_FANGSONG_GB2312)


def _format_cell_runs(cell, size_pt: float, *, bold: bool = False) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            _fmt_run(run, size_pt, bold=bold)


def _set_cell_text(cell, value: str, *, size_pt: float, bold: bool = False) -> None:
    cell.text = str(value or "")
    _format_cell_runs(cell, size_pt, bold=bold)


def _cell_text(table, row: int, col: int, value: str, *, bold: bool = False) -> None:
    _set_cell_text(table.cell(row, col), value, size_pt=PT_XIAO_SAN, bold=bold)


def _normalize_items(payload: dict) -> list[dict[str, str]]:
    raw = payload.get("evidence_items") or payload.get("evidenceItems") or []
    if not isinstance(raw, list):
        return []
    out: list[dict[str, str]] = []
    for it in raw:
        if not isinstance(it, dict):
            continue
        out.append(
            {
                "name": str(it.get("name") or it.get("evidence_name") or "").strip(),
                "source": str(it.get("source") or "").strip(),
                "description": str(it.get("description") or it.get("note") or "").strip(),
                "pages": str(it.get("pages") or "").strip(),
            }
        )
    return out


def _parse_int_loose(text: str) -> int | None:
    t = str(text or "").strip()
    if not t:
        return None
    try:
        return int(float(t))
    except ValueError:
        return None


def build_evidence_list_document(fields: dict) -> Document:
    """根据结构化字段生成证据材料清单文档。"""
    items = _normalize_items(fields)
    submitter = str(fields.get("submitter_name") or fields.get("submitterName") or "").strip()
    sub_date = str(fields.get("submission_date") or fields.get("submissionDate") or "").strip()
    receiver = str(fields.get("court_receiver") or fields.get("courtReceiver") or "").strip()

    total_items_txt = str(fields.get("total_items") or fields.get("totalItems") or "").strip()
    total_pages_txt = str(fields.get("total_pages") or fields.get("totalPages") or "").strip()

    nonempty = [it for it in items if any(str(v).strip() for v in it.values())]
    display_rows = items if items else [{"name": "", "source": "", "description": "", "pages": ""}]

    if total_items_txt:
        try:
            n_items = int(total_items_txt)
        except ValueError:
            n_items = len(nonempty) if nonempty else 0
    else:
        n_items = len(nonempty) if nonempty else 0

    if total_pages_txt:
        total_pages = total_pages_txt
    else:
        s = 0
        ok = False
        for it in nonempty:
            p = _parse_int_loose(it.get("pages") or "")
            if p is not None:
                s += p
                ok = True
        total_pages = str(s) if ok else ""

    doc = Document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("证据材料清单")
    _fmt_run(tr, PT_YI_HAO, bold=True)

    doc.add_paragraph("")

    nrows = 1 + len(display_rows) + 1
    table = doc.add_table(rows=nrows, cols=5)
    table.style = "Table Grid"

    headers = ("序号", "证据名称", "证据来源", "证据说明", "页数")
    for col, h in enumerate(headers):
        _cell_text(table, 0, col, h, bold=True)

    for idx, it in enumerate(display_rows):
        r = idx + 1
        _cell_text(table, r, 0, str(idx + 1))
        _cell_text(table, r, 1, it.get("name") or "")
        _cell_text(table, r, 2, it.get("source") or "")
        _cell_text(table, r, 3, it.get("description") or "")
        _cell_text(table, r, 4, it.get("pages") or "")

    summary_row = table.rows[-1]
    cell0 = summary_row.cells[0]
    cell0.merge(summary_row.cells[4])
    items_display = total_items_txt if total_items_txt else (str(n_items) if n_items else "　")
    pages_display = total_pages_txt if total_pages_txt else (str(total_pages) if total_pages else "　")
    _set_cell_text(
        cell0,
        f"证据合共 {items_display} 项 {pages_display} 页（以上材料均为复印件）",
        size_pt=PT_XIAO_SAN,
        bold=False,
    )

    doc.add_paragraph("")
    line = doc.add_paragraph()
    r_sub = line.add_run(f"提交人：{submitter or '　　'}")
    _fmt_run(r_sub, PT_XIAO_SAN)
    r_tab = line.add_run("\t\t")
    _fmt_run(r_tab, PT_XIAO_SAN)
    r_date = line.add_run(f"日期：{sub_date or '　　　　年　　月　　日'}")
    _fmt_run(r_date, PT_XIAO_SAN)

    p_recv = doc.add_paragraph()
    r_recv = p_recv.add_run(f"法院接收人：{receiver or '　　　'}")
    _fmt_run(r_recv, PT_XIAO_SAN)

    return doc


def evidence_list_docx_bytes(fields: dict) -> bytes:
    doc = build_evidence_list_document(fields)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def ensure_evidence_list_template_exists() -> Path:
    """
    保留实体模板文件（与其他文书一致）。

    证据清单当前仍采用程序生成策略；该模板用于归档和后续人工校验样式基线。
    """
    path = EVIDENCE_LIST_TEMPLATE_PATH
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        return path
    template_doc = build_evidence_list_document(
        {
            "evidence_items": [
                {
                    "name": "［证据名称］",
                    "source": "［证据来源］",
                    "description": "［证据说明］",
                    "pages": "［页数］",
                }
            ],
            "submitter_name": "［提交人］",
            "submission_date": "［　　年　　月　　日］",
            "court_receiver": "［法院接收人］",
            "total_items": "",
            "total_pages": "",
        }
    )
    template_doc.save(path)
    return path
