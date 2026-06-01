import json
import re
from io import BytesIO
from pathlib import Path

from flask import jsonify, request, send_file
from docx import Document

from .civil_complaint_docx import (
    CIVIL_COMPLAINT_TEMPLATE_PATH,
    build_civil_complaint_mapping,
    ensure_civil_complaint_template_exists,
    replace_placeholders_in_document,
)
from .enforcement_application_docx import (
    ENFORCEMENT_APPLICATION_TEMPLATE_PATH,
    build_enforcement_application_mapping,
    ensure_enforcement_application_template_exists,
)
from .evidence_list_docx import (
    EVIDENCE_LIST_TEMPLATE_PATH,
    ensure_evidence_list_template_exists,
    evidence_list_docx_bytes,
)
from .labor_arbitration_docx import (
    LABOR_ARBITRATION_TEMPLATE_PATH,
    build_labor_arbitration_mapping,
    ensure_labor_arbitration_template_exists,
)
from .labor_mediation_docx import (
    LABOR_MEDIATION_APPLICATION_TEMPLATE_PATH,
    build_labor_mediation_mapping,
    ensure_labor_mediation_application_template_exists,
)

DOCUMENT_FIELD_SCHEMA_LABOR_COMPLAINT = [
    {"key": "complainant_name", "label": "姓名"},
    {"key": "complainant_gender", "label": "性别"},
    {"key": "complainant_mobile_phone", "label": "手机号码"},
    {"key": "complainant_id_number", "label": "身份证件号"},
    {"key": "complainant_mailing_address", "label": "通讯地址"},
    {"key": "complainant_landline_phone", "label": "固定电话"},
    {"key": "complainant_postal_code", "label": "邮编"},
    {"key": "respondent_name", "label": "名称(姓名)"},
    {"key": "respondent_legal_representative", "label": "法定代表人(主要负责人)"},
    {"key": "respondent_contact_name", "label": "姓名"},
    {"key": "respondent_contact_job_title", "label": "职务"},
    {"key": "respondent_registered_address", "label": "注册地址"},
    {"key": "respondent_business_address", "label": "实际经营地址"},
    {"key": "respondent_contact_phone", "label": "联系电话"},
    {"key": "respondent_postal_code", "label": "邮编"},
    {"key": "claim_requests", "label": "请求事项"},
    {"key": "facts_and_reasons", "label": "事实与理由"},
]
LABOR_COMPLAINT_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent / "docx" / "laborSecurityInspectionComplaintForm.docx"
)


def _empty_labor_complaint_fields() -> dict:
    return {item["key"]: "" for item in DOCUMENT_FIELD_SCHEMA_LABOR_COMPLAINT}


def _extract_json_object_from_text(text: str) -> dict | None:
    raw = str(text or "").strip()
    if not raw:
        return None
    try:
        parsed = json.loads(raw)
        return parsed if isinstance(parsed, dict) else None
    except json.JSONDecodeError:
        pass

    m = re.search(r"\{[\s\S]*\}", raw)
    if not m:
        return None
    try:
        parsed = json.loads(m.group(0))
        return parsed if isinstance(parsed, dict) else None
    except json.JSONDecodeError:
        return None


def _pick(payload: dict, *keys: str) -> str:
    for key in keys:
        value = payload.get(key)
        if value is None:
            continue
        text = str(value).strip()
        if text:
            return text
    return ""


def _fill_table_cell(table, row: int, col: int, value: str) -> None:
    cell = table.cell(row, col)
    text = str(value or "")
    # python-docx: assigning cell.text clears existing paragraphs/runs reliably.
    cell.text = text


def register_document_writing_routes(app, *, chat_completion, resolve_user_id_from_token) -> None:
    @app.route("/api/documents/polish", methods=["POST", "OPTIONS"])
    def api_documents_polish():
        """文书草稿 LLM 润色（混元）；不改变用户自述事实，仅规范法律用语与结构。"""
        if request.method == "OPTIONS":
            return ("", 204)

        user_id = resolve_user_id_from_token()
        if user_id is None:
            return jsonify({"message": "未登录或令牌无效，请重新登录"}), 401

        payload = request.get_json(silent=True) or {}
        doc_type = str(payload.get("doc_type") or "arbitration").strip()
        raw_text = str(payload.get("raw_text") or "").strip()
        if not raw_text:
            return jsonify({"message": "请先生成或粘贴文书正文后再润色"}), 400
        if len(raw_text) > 50000:
            return jsonify({"message": "正文过长，请删减后再试"}), 400

        type_names = {
            "arbitration": "劳动仲裁申请书",
            "labor_arbitration_application": "劳动人事争议仲裁申请书",
            "labor_mediation_application": "劳动争议调解申请书",
            "complaint": "劳动监察投诉书",
            "civil_complaint": "民事起诉状",
            "enforcement_application": "申请执行书",
            "evidence_list": "证据材料清单",
            "demand": "律师函（催告函）",
        }
        type_cn = type_names.get(doc_type, "劳动争议文书")

        if doc_type == "civil_complaint":
            system = (
                "你是熟悉中国民事诉讼实务的法律文书助手。用户会提供自行填写的民事起诉状草稿。"
                "要求：在不变更已陈述的核心事实、诉讼请求、金额与日期等大前提的前提下，将表述调整为更规范的书面法律用语；"
                "保持当事人信息、案由、此致法院与具状人结构；严禁编造任何未出现的新事实、新证据或具体金额。"
                "不要输出 Markdown 代码围栏，不要加开场白或结束语，直接输出可提交的文书正文。"
            )
        elif doc_type == "enforcement_application":
            system = (
                "你是熟悉中国民事强制执行实务的法律文书助手。用户会提供自行填写的申请执行书草稿。"
                "要求：在不变更已陈述的执行依据（案号、作出机关、生效情况）、申请执行事项、金额与日期等大前提的前提下，"
                "将表述调整为更规范的书面法律用语；保持申请执行人、被申请执行人、执行依据、此致法院与申请人签章结构；"
                "严禁编造任何未出现的新事实、新依据或具体金额。"
                "不要输出 Markdown 代码围栏，不要加开场白或结束语，直接输出可提交的文书正文。"
            )
        elif doc_type == "evidence_list":
            system = (
                "你是熟悉中国民事诉讼材料整理的法律文书助手。用户会提供证据材料清单的文本版（含表格化条目）。"
                "要求：在不变更证据名称、来源、说明、页数及合计等已填内容的前提下，统一用语与标点，使条目表述简洁规范；"
                "不要增删证据条目，不要编造未出现的材料；保持「证据合共…项…页」与提交人、日期、法院接收人结构。"
                "不要输出 Markdown 代码围栏，不要加开场白或结束语，直接输出润色后的全文。"
            )
        elif doc_type in ("labor_arbitration_application", "arbitration"):
            system = (
                "你是熟悉中国劳动人事争议仲裁实务的法律文书助手。用户会提供自行填写的劳动仲裁申请书草稿。"
                "要求：在不变更已陈述的劳动关系事实、仲裁请求、金额与日期等大前提的前提下，将表述调整为更规范的书面法律用语；"
                "保持申请人、被申请人、仲裁请求、事实与理由、证据、此致仲裁委员会与申请人签章结构；"
                "严禁编造任何未出现的新事实、新证据或具体金额。"
                "不要输出 Markdown 代码围栏，不要加开场白或结束语，直接输出可提交的文书正文。"
            )
        elif doc_type == "labor_mediation_application":
            system = (
                "你是熟悉中国劳动争议调解实务的法律文书助手。用户会提供自行填写的劳动争议调解申请书草稿。"
                "要求：在不变更已陈述的争议事实、调解请求、金额与日期等大前提的前提下，将表述调整为更规范的书面法律用语；"
                "保持申请人、被申请人、调解请求、事实与理由、证据、此致调解组织与申请人签章结构；"
                "严禁编造任何未出现的新事实、新证据或具体金额。"
                "不要输出 Markdown 代码围栏，不要加开场白或结束语，直接输出可提交的文书正文。"
            )
        else:
            system = (
                "你是熟悉中国劳动争议实务的法律文书助手。用户会提供自行填写的文书草稿。"
                "要求：在不变更已陈述的核心事实、诉求、金额与日期等大前提的前提下，将表述调整为更规范的书面法律用语；"
                "可适当分段、编号，使结构清晰；严禁编造任何未出现的新事实、新证据或具体金额。"
                "不要输出 Markdown 代码围栏，不要加开场白或结束语，直接输出可提交的文书正文。"
            )
        user_msg = f"文书类型：{type_cn}\n\n以下是用户草稿，请润色：\n\n{raw_text}"

        try:
            text = chat_completion(
                [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_msg},
                ]
            )
        except ValueError as exc:
            return jsonify({"message": str(exc)}), 503
        except Exception as exc:
            return jsonify({"message": f"润色失败：{exc}"}), 500

        return jsonify({"text": (text or "").strip()})

    @app.route("/api/documents/extract-fields", methods=["POST", "OPTIONS"])
    def api_documents_extract_fields():
        """从用户输入中抽取劳动监察投诉书字段，返回结构化英文变量名。"""
        if request.method == "OPTIONS":
            return ("", 204)

        user_id = resolve_user_id_from_token()
        if user_id is None:
            return jsonify({"message": "未登录或令牌无效，请重新登录"}), 401

        payload = request.get_json(silent=True) or {}
        source_text = str(payload.get("source_text") or payload.get("raw_text") or "").strip()
        if not source_text:
            return jsonify({"message": "请提供待抽取的文本内容（source_text）"}), 400
        if len(source_text) > 50000:
            return jsonify({"message": "文本过长，请删减后再试"}), 400

        fields = _empty_labor_complaint_fields()
        schema_lines = "\n".join(
            f'- "{item["key"]}"（{item["label"]}）' for item in DOCUMENT_FIELD_SCHEMA_LABOR_COMPLAINT
        )
        system = (
            "你是法律文书结构化抽取助手。"
            "请从输入文本中抽取劳动监察投诉书字段。"
            "仅返回一个 JSON 对象，不要输出代码围栏、注释或任何额外说明。"
            "所有字段必须存在，值一律为字符串；无法确定时返回空字符串。"
        )
        user_msg = (
            "请按以下字段键名输出 JSON：\n"
            f"{schema_lines}\n\n"
            "待抽取文本如下：\n"
            f"{source_text}"
        )
        try:
            llm_text = chat_completion(
                [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_msg},
                ]
            )
        except ValueError as exc:
            return jsonify({"message": str(exc)}), 503
        except Exception as exc:
            return jsonify({"message": f"字段抽取失败：{exc}"}), 500

        parsed = _extract_json_object_from_text(llm_text)
        if parsed:
            for key in fields.keys():
                value = parsed.get(key, "")
                fields[key] = str(value).strip() if value is not None else ""

        return jsonify(
            {
                "fields": fields,
                "field_schema": DOCUMENT_FIELD_SCHEMA_LABOR_COMPLAINT,
                "raw_model_output": str(llm_text or "").strip(),
            }
        )

    @app.route("/api/documents/labor-complaint-docx", methods=["POST", "OPTIONS"])
    def api_documents_labor_complaint_docx():
        if request.method == "OPTIONS":
            return ("", 204)

        user_id = resolve_user_id_from_token()
        if user_id is None:
            return jsonify({"message": "未登录或令牌无效，请重新登录"}), 401

        if not LABOR_COMPLAINT_TEMPLATE_PATH.exists():
            return jsonify({"message": "劳动保障监察投诉书模板不存在，请检查 backend/docx 目录"}), 500

        payload = request.get_json(silent=True) or {}
        fields = {
            "complainant_name": _pick(payload, "complainant_name", "applicant"),
            "complainant_gender": _pick(payload, "complainant_gender", "complainantGender"),
            "complainant_mobile_phone": _pick(payload, "complainant_mobile_phone", "applicantPhone"),
            "complainant_id_number": _pick(payload, "complainant_id_number", "complainantIdNumber"),
            "complainant_mailing_address": _pick(payload, "complainant_mailing_address", "complainantAddress"),
            "complainant_landline_phone": _pick(payload, "complainant_landline_phone"),
            "complainant_postal_code": _pick(payload, "complainant_postal_code", "complainantPostalCode"),
            "respondent_name": _pick(payload, "respondent_name", "respondent"),
            "respondent_legal_representative": _pick(
                payload, "respondent_legal_representative", "respondentLegalRepresentative"
            ),
            "respondent_contact_name": _pick(payload, "respondent_contact_name", "respondentContactName"),
            "respondent_contact_job_title": _pick(
                payload, "respondent_contact_job_title", "respondentContactJobTitle"
            ),
            "respondent_registered_address": _pick(
                payload, "respondent_registered_address", "respondentRegisteredAddress"
            ),
            "respondent_business_address": _pick(payload, "respondent_business_address", "respondentBusinessAddress"),
            "respondent_contact_phone": _pick(payload, "respondent_contact_phone", "respondentContactPhone"),
            "respondent_postal_code": _pick(payload, "respondent_postal_code", "respondentPostalCode"),
            "claim_requests": _pick(payload, "claim_requests", "claims"),
            "facts_and_reasons": _pick(payload, "facts_and_reasons", "facts"),
        }

        try:
            doc = Document(str(LABOR_COMPLAINT_TEMPLATE_PATH))
            table = doc.tables[0]
            _fill_table_cell(table, 0, 2, fields["complainant_name"])
            _fill_table_cell(table, 0, 4, fields["complainant_gender"])
            _fill_table_cell(table, 0, 6, fields["complainant_mobile_phone"])
            _fill_table_cell(table, 1, 2, fields["complainant_id_number"])
            _fill_table_cell(table, 1, 6, fields["complainant_landline_phone"])
            _fill_table_cell(table, 2, 2, fields["complainant_mailing_address"])
            _fill_table_cell(table, 2, 6, fields["complainant_postal_code"])
            _fill_table_cell(table, 3, 2, fields["respondent_name"])
            _fill_table_cell(table, 3, 6, fields["respondent_registered_address"])
            _fill_table_cell(table, 4, 6, fields["respondent_business_address"])
            _fill_table_cell(table, 5, 2, fields["respondent_legal_representative"])
            _fill_table_cell(table, 6, 2, fields["respondent_contact_name"])
            _fill_table_cell(table, 5, 4, fields["respondent_contact_job_title"])
            _fill_table_cell(table, 6, 4, fields["respondent_contact_job_title"])
            _fill_table_cell(table, 5, 6, fields["respondent_contact_phone"])
            _fill_table_cell(table, 6, 6, fields["respondent_postal_code"])
            _fill_table_cell(table, 7, 0, fields["claim_requests"])
            _fill_table_cell(table, 8, 0, fields["facts_and_reasons"])
        except Exception as exc:
            return jsonify({"message": f"填充 Word 模板失败：{exc}"}), 500

        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return send_file(
            out,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="劳动保障监察投诉书_已填写.docx",
        )

    @app.route("/api/documents/civil-complaint-docx", methods=["POST", "OPTIONS"])
    def api_documents_civil_complaint_docx():
        if request.method == "OPTIONS":
            return ("", 204)

        user_id = resolve_user_id_from_token()
        if user_id is None:
            return jsonify({"message": "未登录或令牌无效，请重新登录"}), 401

        try:
            ensure_civil_complaint_template_exists()
        except Exception as exc:
            return jsonify({"message": f"无法准备民事起诉状模板：{exc}"}), 500

        if not CIVIL_COMPLAINT_TEMPLATE_PATH.exists():
            return jsonify({"message": "民事起诉状模板不存在，请检查 backend/docx 目录"}), 500

        payload = request.get_json(silent=True) or {}
        fields = {
            "plaintiff_name": _pick(payload, "plaintiff_name", "plaintiffName"),
            "plaintiff_gender": _pick(payload, "plaintiff_gender", "plaintiffGender"),
            "plaintiff_ethnicity": _pick(payload, "plaintiff_ethnicity", "plaintiffEthnicity"),
            "plaintiff_birth": _pick(payload, "plaintiff_birth", "plaintiffBirth"),
            "plaintiff_address": _pick(payload, "plaintiff_address", "plaintiffAddress"),
            "plaintiff_id_number": _pick(payload, "plaintiff_id_number", "plaintiffIdNumber"),
            "plaintiff_phone": _pick(payload, "plaintiff_phone", "plaintiffPhone"),
            "defendant_name": _pick(payload, "defendant_name", "defendantName"),
            "defendant_address": _pick(payload, "defendant_address", "defendantAddress"),
            "defendant_phone": _pick(payload, "defendant_phone", "defendantPhone"),
            "defendant_legal_representative": _pick(
                payload, "defendant_legal_representative", "defendantLegalRepresentative"
            ),
            "case_cause": _pick(payload, "case_cause", "caseCause"),
            "claims": _pick(payload, "claims"),
            "facts": _pick(payload, "facts"),
            "evidence_list": _pick(payload, "evidence_list", "evidenceList"),
            "court_name": _pick(payload, "court_name", "courtName"),
            "date_text": _pick(payload, "date_text", "dateText"),
        }

        mapping = build_civil_complaint_mapping(fields)

        try:
            doc = Document(str(CIVIL_COMPLAINT_TEMPLATE_PATH))
            replace_placeholders_in_document(doc, mapping)
        except Exception as exc:
            return jsonify({"message": f"填充 Word 模板失败：{exc}"}), 500

        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return send_file(
            out,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="民事起诉状_已填写.docx",
        )

    @app.route("/api/documents/enforcement-application-docx", methods=["POST", "OPTIONS"])
    def api_documents_enforcement_application_docx():
        if request.method == "OPTIONS":
            return ("", 204)

        user_id = resolve_user_id_from_token()
        if user_id is None:
            return jsonify({"message": "未登录或令牌无效，请重新登录"}), 401

        try:
            ensure_enforcement_application_template_exists()
        except Exception as exc:
            return jsonify({"message": f"无法准备申请执行书模板：{exc}"}), 500

        if not ENFORCEMENT_APPLICATION_TEMPLATE_PATH.exists():
            return jsonify({"message": "申请执行书模板不存在，请检查 backend/docx 目录"}), 500

        payload = request.get_json(silent=True) or {}
        fields = {
            "applicant_name": _pick(payload, "applicant_name", "applicantName"),
            "applicant_gender": _pick(payload, "applicant_gender", "applicantGender"),
            "applicant_ethnicity": _pick(payload, "applicant_ethnicity", "applicantEthnicity"),
            "applicant_birth": _pick(payload, "applicant_birth", "applicantBirth"),
            "applicant_job": _pick(payload, "applicant_job", "applicantJob"),
            "applicant_address": _pick(payload, "applicant_address", "applicantAddress"),
            "applicant_id_number": _pick(payload, "applicant_id_number", "applicantIdNumber"),
            "applicant_phone": _pick(payload, "applicant_phone", "applicantPhone"),
            "respondent_name": _pick(payload, "respondent_name", "respondentName"),
            "respondent_address": _pick(payload, "respondent_address", "respondentAddress"),
            "respondent_phone": _pick(payload, "respondent_phone", "respondentPhone"),
            "respondent_legal_representative": _pick(
                payload, "respondent_legal_representative", "respondentLegalRepresentative"
            ),
            "legal_representative_line": _pick(
                payload, "legal_representative_line", "legalRepresentativeLine"
            ),
            "entrusted_agent_line": _pick(payload, "entrusted_agent_line", "entrustedAgentLine"),
            "case_cause": _pick(payload, "case_cause", "caseCause"),
            "basis_judgment_no": _pick(payload, "basis_judgment_no", "basisJudgmentNo"),
            "basis_issuer": _pick(payload, "basis_issuer", "basisIssuer"),
            "basis_effective_date": _pick(payload, "basis_effective_date", "basisEffectiveDate"),
            "basis_extra": _pick(payload, "basis_extra", "basisExtra"),
            "basis_doc_type_phrase": _pick(payload, "basis_doc_type_phrase", "basisDocTypePhrase"),
            "enforcement_non_performance_phrase": _pick(
                payload, "enforcement_non_performance_phrase", "enforcementNonPerformancePhrase"
            ),
            "requests": _pick(payload, "requests", "enforcementRequests"),
            "facts": _pick(payload, "facts"),
            "court_name": _pick(payload, "court_name", "courtName"),
            "attachment_line": _pick(payload, "attachment_line", "attachmentLine"),
            "date_text": _pick(payload, "date_text", "dateText"),
        }

        mapping = build_enforcement_application_mapping(fields)

        try:
            doc = Document(str(ENFORCEMENT_APPLICATION_TEMPLATE_PATH))
            replace_placeholders_in_document(doc, mapping)
        except Exception as exc:
            return jsonify({"message": f"填充 Word 模板失败：{exc}"}), 500

        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return send_file(
            out,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="申请执行书_已填写.docx",
        )

    @app.route("/api/documents/evidence-list-docx", methods=["POST", "OPTIONS"])
    def api_documents_evidence_list_docx():
        if request.method == "OPTIONS":
            return ("", 204)

        user_id = resolve_user_id_from_token()
        if user_id is None:
            return jsonify({"message": "未登录或令牌无效，请重新登录"}), 401

        try:
            ensure_evidence_list_template_exists()
        except Exception as exc:
            return jsonify({"message": f"无法准备证据材料清单模板：{exc}"}), 500

        if not EVIDENCE_LIST_TEMPLATE_PATH.exists():
            return jsonify({"message": "证据材料清单模板不存在，请检查 backend/docx 目录"}), 500

        payload = request.get_json(silent=True) or {}
        raw_items = payload.get("evidence_items")
        if raw_items is None:
            raw_items = payload.get("evidenceItems")
        if not isinstance(raw_items, list):
            raw_items = []

        fields = {
            "evidence_items": raw_items,
            "submitter_name": _pick(payload, "submitter_name", "submitterName"),
            "submission_date": _pick(payload, "submission_date", "submissionDate"),
            "court_receiver": _pick(payload, "court_receiver", "courtReceiver"),
            "total_items": _pick(payload, "total_items", "totalItems"),
            "total_pages": _pick(payload, "total_pages", "totalPages"),
        }

        try:
            data = evidence_list_docx_bytes(fields)
        except Exception as exc:
            return jsonify({"message": f"生成证据材料清单 Word 失败：{exc}"}), 500

        out = BytesIO(data)
        out.seek(0)
        return send_file(
            out,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="证据材料清单_已填写.docx",
        )

    @app.route("/api/documents/labor-arbitration-application-docx", methods=["POST", "OPTIONS"])
    def api_documents_labor_arbitration_application_docx():
        if request.method == "OPTIONS":
            return ("", 204)

        user_id = resolve_user_id_from_token()
        if user_id is None:
            return jsonify({"message": "未登录或令牌无效，请重新登录"}), 401

        try:
            ensure_labor_arbitration_template_exists()
        except Exception as exc:
            return jsonify({"message": f"无法准备劳动仲裁申请书模板：{exc}"}), 500

        if not LABOR_ARBITRATION_TEMPLATE_PATH.exists():
            return jsonify({"message": "劳动仲裁申请书模板不存在，请检查 backend/docx 目录"}), 500

        payload = request.get_json(silent=True) or {}
        fields = {
            "applicant_name": _pick(payload, "applicant_name", "applicantName"),
            "applicant_gender": _pick(payload, "applicant_gender", "applicantGender"),
            "applicant_ethnicity": _pick(payload, "applicant_ethnicity", "applicantEthnicity"),
            "applicant_birth": _pick(payload, "applicant_birth", "applicantBirth"),
            "applicant_address": _pick(payload, "applicant_address", "applicantAddress"),
            "applicant_id_type": _pick(payload, "applicant_id_type", "applicantIdType"),
            "applicant_id_number": _pick(payload, "applicant_id_number", "applicantIdNumber"),
            "applicant_job": _pick(payload, "applicant_job", "applicantJob"),
            "applicant_phone": _pick(payload, "applicant_phone", "applicantPhone"),
            "contract_performance_place": _pick(
                payload, "contract_performance_place", "contractPerformancePlace"
            ),
            "respondent_name": _pick(payload, "respondent_name", "respondentName"),
            "respondent_address": _pick(payload, "respondent_address", "respondentAddress"),
            "respondent_phone": _pick(payload, "respondent_phone", "respondentPhone"),
            "respondent_legal_representative": _pick(
                payload, "respondent_legal_representative", "respondentLegalRepresentative"
            ),
            "respondent_legal_representative_job": _pick(
                payload, "respondent_legal_representative_job", "respondentLegalRepresentativeJob"
            ),
            "respondent_business_place": _pick(
                payload, "respondent_business_place", "respondentBusinessPlace"
            ),
            "respondent_contact_person": _pick(
                payload, "respondent_contact_person", "respondentContactPerson"
            ),
            "claims": _pick(payload, "claims"),
            "facts": _pick(payload, "facts"),
            "evidence_list": _pick(payload, "evidence_list", "evidenceList"),
            "arbitration_commission": _pick(payload, "arbitration_commission", "arbitrationCommission"),
            "agent_block": _pick(payload, "agent_block", "agentBlock"),
            "attachment_line": _pick(payload, "attachment_line", "attachmentLine"),
            "date_text": _pick(payload, "date_text", "dateText"),
        }

        mapping = build_labor_arbitration_mapping(fields)

        try:
            doc = Document(str(LABOR_ARBITRATION_TEMPLATE_PATH))
            replace_placeholders_in_document(doc, mapping)
        except Exception as exc:
            return jsonify({"message": f"填充 Word 模板失败：{exc}"}), 500

        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return send_file(
            out,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="劳动人事争议仲裁申请书_已填写.docx",
        )

    @app.route("/api/documents/labor-mediation-application-docx", methods=["POST", "OPTIONS"])
    def api_documents_labor_mediation_application_docx():
        if request.method == "OPTIONS":
            return ("", 204)

        user_id = resolve_user_id_from_token()
        if user_id is None:
            return jsonify({"message": "未登录或令牌无效，请重新登录"}), 401

        try:
            ensure_labor_mediation_application_template_exists()
        except Exception as exc:
            return jsonify({"message": f"无法准备劳动争议调解申请书模板：{exc}"}), 500

        if not LABOR_MEDIATION_APPLICATION_TEMPLATE_PATH.exists():
            return jsonify({"message": "劳动争议调解申请书模板不存在，请检查 backend/docx 目录"}), 500

        payload = request.get_json(silent=True) or {}
        fields = {
            "applicant_name": _pick(payload, "applicant_name", "applicantName"),
            "applicant_gender": _pick(payload, "applicant_gender", "applicantGender"),
            "applicant_ethnicity": _pick(payload, "applicant_ethnicity", "applicantEthnicity"),
            "applicant_birth": _pick(payload, "applicant_birth", "applicantBirth"),
            "applicant_address": _pick(payload, "applicant_address", "applicantAddress"),
            "applicant_id_number": _pick(payload, "applicant_id_number", "applicantIdNumber"),
            "applicant_phone": _pick(payload, "applicant_phone", "applicantPhone"),
            "respondent_name": _pick(payload, "respondent_name", "respondentName"),
            "respondent_address": _pick(payload, "respondent_address", "respondentAddress"),
            "respondent_phone": _pick(payload, "respondent_phone", "respondentPhone"),
            "respondent_legal_representative": _pick(
                payload, "respondent_legal_representative", "respondentLegalRepresentative"
            ),
            "claims": _pick(payload, "claims"),
            "facts": _pick(payload, "facts"),
            "evidence_list": _pick(payload, "evidence_list", "evidenceList"),
            "mediation_org": _pick(payload, "mediation_org", "mediationOrg"),
            "mediation_preamble": _pick(payload, "mediation_preamble", "mediationPreamble"),
            "date_text": _pick(payload, "date_text", "dateText"),
        }

        mapping = build_labor_mediation_mapping(fields)

        try:
            doc = Document(str(LABOR_MEDIATION_APPLICATION_TEMPLATE_PATH))
            replace_placeholders_in_document(doc, mapping)
        except Exception as exc:
            return jsonify({"message": f"填充 Word 模板失败：{exc}"}), 500

        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return send_file(
            out,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="劳动争议调解申请书_已填写.docx",
        )
